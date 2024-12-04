import os
from django.conf import settings

from django.shortcuts import render, redirect
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import Document
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2

from pymongo import MongoClient
from .forms import ArchivoForm
import datetime
from .models import Archivo  # Importa tu modelo
from django.contrib import messages
from datetime import datetime
from django.utils.html import format_html

from django.core.paginator import Paginator

from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from werkzeug.security import generate_password_hash, check_password_hash
from django.contrib.auth.models import User
from .models import Archivo



# Conexión a MongoDB 
client = MongoClient('mongodb://localhost:27017/') 
db = client['plagioInspector'] 
collection = db['archivos']


# Vista de login
def login_view(request):

    if request.user.is_authenticated:  # Si el usuario ya está autenticado
        return redirect('index')       # Redirige al index

    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')
        
        if not email or not password:
            return render(request, 'pc/login.html', {
                'error': 'Por favor, completa todos los campos.'
            })

        user = authenticate(request, username=email, password=password)
        if user is not None:
            login(request, user)
            next_url = request.GET.get('next', 'index')  # Redirige a la URL solicitada o al index
            return redirect(next_url)
        else:
            return render(request, 'pc/login.html', {
                'error': 'Credenciales inválidas. Por favor, inténtalo de nuevo.'
            })

    return render(request, 'pc/login.html')


# Registro de usuario
from django.http import JsonResponse


def register(request):
    if request.method == 'POST':
        email = request.POST['email']
        password = request.POST['password']
        if User.objects.filter(email=email).exists():
            context = {
                'message': 'El correo ya ha sido registrado.',
                'success': False,
            }
        else:
            # Lógica para crear el usuario
            User.objects.create_user(username=email, email=email, password=password)
            context = {
                'message': 'Registro exitoso. ¡Ahora puedes iniciar sesión!',
                'success': True,
            }
        return render(request, 'pc/register.html', context)
    return render(request, 'pc/register.html')


# Home
def home(request):
    return render(request, 'pc/login.html')

def fileAyuda(request):
    return render(request, 'pc/ayuda.html')

@login_required
def index_view(request): 
    usuario = request.user  # Usuario autenticado
    
    # Obtiene solo los archivos del usuario logueado
    archivos = Archivo.objects.filter(usuario=request.user)
    context = {'archivos': archivos}
    return render(request, 'pc/index.html', context)


@login_required
def fileCompare(request): 
    if not request.user.is_authenticated:
        return redirect('/login/') 

    archivos = Archivo.objects.filter(usuario=request.user)
    context = {'archivos': archivos}
    return render(request, 'pc/doc_compare.html', context)

# New Vista Historial
@login_required
def fileHistorial(request):
    if not request.user.is_authenticated:
        return redirect('/login/') 

    archivos = Archivo.objects.filter(usuario=request.user)
    context = {'archivos': archivos}
    return render(request, 'pc/historial.html', context)


# New vista para revisar un documento dentro del historial
@login_required 
def fileRevisar(request): 
    user_files = File.objects.filter(user=request.user)
    return render(request, 'pc/revisar.html', {'user_files': user_files})


# Web search (Text)
def test(request):
    print("Request received in test view")
    texto = request.POST.get('q', '')
    print("Texto ingresado:", texto)

    if texto:
        percent, link, plagiarized_texts, coincidencias, _ = main.findSimilarity(texto)
        percent = round(percent, 2)

        # Crear oraciones resaltadas
        oraciones = texto.split('.')
        oraciones_resaltadas = []
        for fragmento in oraciones:
            encontrado = False
            for coincidencia in coincidencias:
                if coincidencia['texto'].lower() in fragmento.lower():
                    oraciones_resaltadas.append({
                        'texto': fragmento,
                        'coincidencia': coincidencia['url']
                    })
                    encontrado = True
                    break
            if not encontrado:
                oraciones_resaltadas.append({'texto': fragmento, 'coincidencia': None})
    else:
        percent, link, plagiarized_texts, coincidencias, oraciones_resaltadas = 0, [], [], [], []

    return render(request, 'pc/index.html', {
        'link': link,
        'percent': percent,
        'plagiarized_texts': plagiarized_texts,
        'oraciones_resaltadas': oraciones_resaltadas
    })


# Web search file (.txt, .docx, .pdf)
def filetest(request):
    value = ''
    docfile = request.FILES.get('docfile')  # Obtiene el archivo subido
    
    if docfile:
        if docfile.name.endswith(".txt"):
            value = docfile.read().decode('utf-8')  # Leer el archivo .txt en memoria

        elif docfile.name.endswith(".docx"):
            document = Document(docfile)
            for para in document.paragraphs:
                value += para.text

        elif docfile.name.endswith(".pdf"):
            pdfReader = PyPDF2.PdfReader(docfile)  # Leer el archivo PDF en memoria
            for page_num in range(len(pdfReader.pages)):
                page = pdfReader.pages[page_num]
                value += page.extract_text()

    # Procesa el contenido del archivo (por ejemplo, buscar similitudes)
    percent, link = main.findSimilarity(value)
    percent = round(percent, 2)

    print("Output...................!!!!!!!!", percent, link)
    return render(request, 'pc/index.html', {'link': link, 'percent': percent})

# Two text compare (Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'], request.POST['q2'])
    result = round(result, 2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})

# Two text compare (.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document1 = Document(request.FILES['docfile1'])
        for para in document1.paragraphs:
            value1 += para.text
        document2 = Document(request.FILES['docfile2'])
        for para in document2.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1, value2)
    
    print("Output..................!!!!!!!!", result)
    return render(request, 'pc/doc_compare.html', {'result': result})

@login_required
def fileCompare(request):
    # Obtener el usuario actual
    usuario_actual = request.user.username

    # Obtener la consulta de búsqueda
    query = request.GET.get('q')    
    
    try:
        if query and isinstance(query, str):  # Verifica que query sea una cadena
            archivos = list(
                collection.find({
                    "usuario": usuario_actual,
                    "NOMBRE": {"$regex": f".*{query}.*", "$options": "i"}  # Busca coincidencias parciales
                })
            )
        else:
            # Si no hay consulta, obtiene todos los archivos del usuario
            archivos = list(collection.find({"usuario": usuario_actual}))
    except Exception as e:
        archivos = []
        print(f"Error al consultar los archivos para el usuario {usuario_actual}: {e}")

    plagio_result = request.session.get('plagio_result', None)

    # Configura el paginador, 5 documentos por página
    paginator = Paginator(archivos, 5)
    page_number = request.GET.get('page')  # Obtén el número de la página actual
    page_obj = paginator.get_page(page_number)  # Obtén la página actual

    return render(request, 'pc/doc_compare.html', {
        'archivos': page_obj.object_list,
        'page_obj': page_obj, 
        'plagio_result': plagio_result
    })

# Visualizacion en la vista de historial de los documentos en la bd 
@login_required
def fileHistorial(request): 
    # Obtener el usuario actual
    usuario_actual = request.user.username

    # Obtener la consulta de búsqueda
    query = request.GET.get('q')

    try:
        if query and isinstance(query, str):  # Verifica que query sea una cadena válida
            archivos = list(
                collection.find({
                    "usuario": usuario_actual,
                    "NOMBRE": {"$regex": f".*{query}.*", "$options": "i"}  # Busca coincidencias parciales
                })
            )
        else:
            # Si no hay búsqueda, filtra todos los archivos del usuario actual
            archivos = list(collection.find({"usuario": usuario_actual}))
    except Exception as e:
        archivos = []
        print(f"Error al consultar los archivos para el usuario {usuario_actual}: {e}")

    # Configurar la paginación (5 documentos por página)
    paginator = Paginator(archivos, 5)
    page_number = request.GET.get('page')  # Número de página actual
    page_obj = paginator.get_page(page_number)  # Página actual

    return render(request, 'pc/historial.html', {
        'archivos': page_obj.object_list,
        'page_obj': page_obj
    })


# Parte del historial elimincacion del archivo 
def eliminar_archivo(request, nombre_archivo): 
    if request.method == 'POST': 
        # Buscar el archivo en la bd 
        archivo = collection.find_one({'NOMBRE': nombre_archivo})
        if archivo: 
            # Obtener la ruta del archivo
            ruta_archivo = archivo.get('RUTA')  
            # Eliminar el archivo físicamente si existe en la carpeta Media
            if ruta_archivo and os.path.exists(ruta_archivo):
                try:
                    os.remove(ruta_archivo)  # Eliminar el archivo del sistema de archivos
                except Exception as e:
                    return HttpResponse(f"Error al eliminar el archivo físico: {str(e)}", status=500)
        result = collection.delete_one({'NOMBRE': nombre_archivo}) 
        if result.deleted_count > 0: 
            return redirect('historial') # Redirige a la lista de archivos después de eliminar 
        else: 
            return HttpResponse('Archivo no encontrado.', status=404) 
    return HttpResponse('Método no permitido.', status=405)


# Codigo para subir archivos en la vista doc_compare 
@login_required
def subir_archivo(request): 
    if request.method == 'POST' and request.FILES.get('archivo'): 
        archivo = request.FILES['archivo'] 
        nombre_archivo = archivo.name 
        fecha_subida = datetime.now().isoformat() 

        # Crear un directorio único por usuario en MEDIA_ROOT
        user_directory = os.path.join(settings.MEDIA_ROOT, f"archivos/{request.user.username}")
        os.makedirs(user_directory, exist_ok=True)  # Crea la carpeta si no existe

        # Guardar el archivo en la carpeta del usuario
        file_path = os.path.join(user_directory, nombre_archivo)
        with open(file_path, 'wb+') as destination: 
            for chunk in archivo.chunks(): 
                destination.write(chunk)
        
        # Guardar el archivo en la colección (base de datos)
        try:
            collection.insert_one({ 
                'usuario': request.user.username,  # Relacionar con el usuario autenticado
                'NOMBRE': nombre_archivo, 
                'FECHA_SUBIDA': fecha_subida, 
                'RUTA': file_path
            }) 
        except Exception as e:
            print(f"Error al guardar el archivo en la base de datos: {e}")
            return HttpResponse('Error al guardar el archivo.', status=500)

        
        return redirect('compare') 
    return HttpResponse('Método no permitido.', status=405)


# Web search file (.txt, .docx, .pdf)
def ejecutar_busqueda(request, nombre_archivo):
    value = ''
    archivo = collection.find_one({'NOMBRE': nombre_archivo})

    docfile_path = archivo['RUTA'] # Obtener el archivo del nombre
    if not docfile_path: 
        return HttpResponse('Ruta del archivo no encontrada en la base de datos.', status=404)

    # lee el contenido del archivo
    if docfile_path.endswith(".txt"): 
        try: 
            with open(docfile_path, 'r', encoding='utf-8') as file: 
                value = file.read() # Leer el archivo .txt
        except Exception as e: 
            return HttpResponse(f"Error al leer el archivo: {str(e)}", status=500)

    elif docfile_path.endswith(".docx"): 
        document = Document(docfile_path) 
        for para in document.paragraphs: 
            value += para.text

    elif docfile_path.endswith(".pdf"): 
        pdfReader = PyPDF2.PdfReader(docfile_path) # Leer el archivo PDF 
        for page_num in range(len(pdfReader.pages)): 
            page = pdfReader.pages[page_num] 
            value += page.extract_text()
        
    # Procesa el contenido del archivo (por ejemplo, buscar similitudes) 
    percent, link, textos_plagiados, _, posiciones  = main.findSimilarity(value) #Se ajusto la funcion para retornar el texto plagiado
    percent = round(percent, 2)

    # Resaltar texto plagiado
    texto_resaltado = value
    offset = 0  # Para ajustar las posiciones al insertar HTML
    for start, end in posiciones:
        start += offset
        end += offset
        texto_resaltado = (
            texto_resaltado[:start] +
            f'<span style="background-color: yellow;">{texto_resaltado[start:end]}</span>' +
            texto_resaltado[end:]
        )
        offset += len('<span style="background-color: yellow;"></span>')


    collection.update_one(
        {'NOMBRE': nombre_archivo},
        {'$set': {
            'PLAGIO': {
                'PORCENTAJE': percent, 
                'ENLACES': link, 
                'TEXTOS_PLAGIADOS': textos_plagiados,
                'TEXTO_RESALTADO': texto_resaltado    # Guardar el texto resaltado
                }
        }}
    )

    print("Output...................!!!!!!!!", percent, link) 
    return redirect('historial')


# funcion leer los archivos 
def leer_archivo(ruta):
    contenido = ''
    if ruta.endswith(".txt"):
        with open(ruta, 'r', encoding='utf-8') as file:
            contenido = file.read()
    elif ruta.endswith(".docx"):
        document = Document(ruta)
        for para in document.paragraphs:
            contenido += para.text
    elif ruta.endswith(".pdf"):
        pdfReader = PyPDF2.PdfReader(ruta)
        for page in pdfReader.pages:
            contenido += page.extract_text()
    return contenido


# New View para comparar entre documentos 
def comparar_documentos(request):
    if request.method == 'POST':
        archivo1_nombre = request.POST.get('archivo1')
        archivo2_nombre = request.POST.get('archivo2')

        # Obtener los archivos de la base de datos
        archivo1 = collection.find_one({'NOMBRE': archivo1_nombre})
        archivo2 = collection.find_one({'NOMBRE': archivo2_nombre})

        # Leer los contenidos de ambos archivos
        contenido1 = leer_archivo(archivo1['RUTA'])
        contenido2 = leer_archivo(archivo2['RUTA'])

        # Calcular la similitud
        percent, sources = main.findSimilarity2(contenido1, contenido2)

        # Renderizar los resultados en una nueva plantilla
        return render(request, 'pc/compare_two.html', {
            'archivo1': archivo1_nombre,
            'archivo2': archivo2_nombre,
            'percent': percent,
            'sources': sources
        })
    return HttpResponse('Método no permitido.', status=405)



# Variable para obtener los datos del archivo plagiado 
def ver_resultados(request, nombre_archivo):
    archivo = collection.find_one({'NOMBRE': nombre_archivo})

    if not archivo or 'PLAGIO' not in archivo:
        return HttpResponse('Resultados de plagio no encontrados.', status=404)

    plagio = archivo.get('PLAGIO', {})
    plagio['TEXTOS_PLAGIADOS'] = plagio.get('TEXTOS_PLAGIADOS', [])

    plagio = archivo['PLAGIO']
    return render(request, 'pc/revisar.html', {'archivo': archivo, 'plagio': plagio})
